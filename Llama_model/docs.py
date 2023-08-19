import io
from docx import Document
from docx.shared import Inches

def create_word_docs(user_input, paragraph, image_input):
    doc = Document()

    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    doc.add_heading('Image', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_input.seek(0)
    doc.add_picture(image_stream, width=Inches(4))

    return doc